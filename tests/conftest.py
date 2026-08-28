from pathlib import Path

import pytest
from docx import Document as DocxDocument

FIXTURE_PARAGRAPHS = [
    ("heading1", "Quarterly Engineering Report"),
    ("heading2", "Introduction"),
    (
        "body",
        "This document summarizes the engineering work completed during the third quarter "
        "of the fiscal year. It covers infrastructure changes, product delivery, and the "
        "operational incidents that shaped our priorities.",
    ),
    (
        "body",
        "Throughout the quarter the team maintained a steady release cadence while absorbing "
        "a significant migration effort. We want to be clear that this was not without cost, "
        "and several planned features slipped as a direct result of that migration.",
    ),
    ("heading2", "Infrastructure"),
    (
        "body",
        "We migrated the primary datastore from a single Postgres instance to a replicated "
        "cluster with automated failover.",
    ),
]


def write_sample_docx(path: Path) -> Path:
    doc = DocxDocument()
    for kind, text in FIXTURE_PARAGRAPHS:
        if kind == "heading1":
            doc.add_heading(text, level=1)
        elif kind == "heading2":
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text)
    doc.save(path)
    return path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """A fresh .docx per test — SuperDoc mutates in place, so tests must not share one."""
    return write_sample_docx(tmp_path / "sample.docx")
